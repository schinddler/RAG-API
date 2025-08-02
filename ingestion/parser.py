"""
Production-grade file parser for RAG backend.

This module handles parsing of various file formats (PDF, DOCX, TXT, EML)
with robust error handling, async support, OCR fallback, MIME validation,
and structured output.
"""

import os
import tempfile
import urllib.request
import urllib.parse
import urllib.error
from pathlib import Path
from typing import Dict, Optional, Union, Tuple, List, Any
import logging
import mimetypes
import email
from email import policy
from email.parser import BytesParser
import asyncio
import aiofiles
import aiohttp
import re
import uuid
from datetime import datetime

# Document parsing libraries
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph

# Enhanced parsing libraries
try:
    import magic
    MAGIC_AVAILABLE = True
except ImportError:
    MAGIC_AVAILABLE = False
    import mimetypes as fallback_mimetypes

try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

try:
    import html2text
    HTML2TEXT_AVAILABLE = True
except ImportError:
    HTML2TEXT_AVAILABLE = False

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class ParserError(Exception):
    """Custom exception for parser-related errors."""
    pass


class FileParser:
    """
    Production-grade file parser for RAG backend.
    
    Supports PDF, DOCX, TXT, and EML files with robust error handling,
    OCR fallback, MIME validation, and async capabilities.
    """
    
    def __init__(self, temp_dir: Optional[str] = None, enable_ocr: bool = False):
        """
        Initialize the file parser.
        
        Args:
            temp_dir: Directory for temporary files. Defaults to system temp.
            enable_ocr: Whether to enable OCR fallback for scanned PDFs.
        """
        self.temp_dir = temp_dir or tempfile.gettempdir()
        self.supported_extensions = {'.pdf', '.docx', '.txt', '.eml'}
        self.max_file_size = 50 * 1024 * 1024  # 50MB limit
        self.enable_ocr = enable_ocr and OCR_AVAILABLE
        
        # MIME type mappings
        self.mime_mappings = {
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.txt': 'text/plain',
            '.eml': 'message/rfc822'
        }
        
    async def parse_from_url(self, url: str) -> Dict[str, Any]:
        """
        Parse a file from a URL.
        
        Args:
            url: The URL pointing to the file to parse.
            
        Returns:
            Dictionary containing parsed text and metadata.
            
        Raises:
            ParserError: If parsing fails or file type is unsupported.
        """
        try:
            # Download file to temporary location
            temp_path = await self._download_file(url)
            
            # Parse the file
            result = await self._parse_file(temp_path, url)
            
            # Clean up temporary file
            self._cleanup_temp_file(temp_path)
            
            return result
            
        except Exception as e:
            logger.error(f"Failed to parse file from URL {url}: {str(e)}")
            raise ParserError(f"Failed to parse file: {str(e)}")
    
    def parse_from_file(self, file_path: str) -> Dict[str, Any]:
        """
        Parse a file from local path.
        
        Args:
            file_path: Path to the local file to parse.
            
        Returns:
            Dictionary containing parsed text and metadata.
            
        Raises:
            ParserError: If parsing fails or file type is unsupported.
        """
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                raise ParserError(f"File not found: {file_path}")
            
            # Validate file size
            file_size = file_path.stat().st_size
            if file_size > self.max_file_size:
                raise ParserError(f"File too large: {file_size} bytes (max: {self.max_file_size})")
            
            # Parse the file
            return self._parse_file_sync(file_path, str(file_path))
            
        except Exception as e:
            logger.error(f"Failed to parse file {file_path}: {str(e)}")
            raise ParserError(f"Failed to parse file: {str(e)}")
    
    async def _download_file(self, url: str) -> str:
        """
        Download a file from URL to temporary location.
        
        Args:
            url: The URL to download from.
            
        Returns:
            Path to the downloaded temporary file.
            
        Raises:
            ParserError: If download fails.
        """
        try:
            # Create temporary file
            temp_file = tempfile.NamedTemporaryFile(
                delete=False, 
                dir=self.temp_dir,
                suffix=self._get_extension_from_url(url)
            )
            temp_path = temp_file.name
            temp_file.close()
            
            # Download file
            async with aiohttp.ClientSession() as session:
                async with session.get(url) as response:
                    if response.status != 200:
                        raise ParserError(f"HTTP {response.status}: {response.reason}")
                    
                    # Check content length
                    content_length = response.headers.get('content-length')
                    if content_length and int(content_length) > self.max_file_size:
                        raise ParserError(f"File too large: {content_length} bytes")
                    
                    # Download content
                    content = await response.read()
                    if len(content) > self.max_file_size:
                        raise ParserError(f"File too large: {len(content)} bytes")
                    
                    # Write to temporary file
                    async with aiofiles.open(temp_path, 'wb') as f:
                        await f.write(content)
            
            logger.info(f"Downloaded file from {url} to {temp_path}")
            return temp_path
            
        except Exception as e:
            # Clean up on error
            if 'temp_path' in locals():
                self._cleanup_temp_file(temp_path)
            raise ParserError(f"Download failed: {str(e)}")
    
    def _get_extension_from_url(self, url: str) -> str:
        """
        Extract file extension from URL.
        
        Args:
            url: The URL to extract extension from.
            
        Returns:
            File extension with leading dot.
        """
        parsed = urllib.parse.urlparse(url)
        path = parsed.path.lower()
        
        # Try to get extension from path
        for ext in self.supported_extensions:
            if path.endswith(ext):
                return ext
        
        # Default to .txt if no extension found
        return '.txt'
    
    def _validate_mime_type(self, file_path: Path, expected_extension: str) -> bool:
        """
        Validate MIME type of file against expected extension.
        
        Args:
            file_path: Path to the file to validate.
            expected_extension: Expected file extension.
            
        Returns:
            True if MIME type matches expected type.
        """
        try:
            if MAGIC_AVAILABLE:
                # Use python-magic for more accurate MIME detection
                mime_type = magic.from_file(str(file_path), mime=True)
            else:
                # Fallback to mimetypes
                mime_type, _ = fallback_mimetypes.guess_type(str(file_path))
                if not mime_type:
                    return True  # Skip validation if we can't determine MIME type
            
            expected_mime = self.mime_mappings.get(expected_extension)
            if expected_mime and mime_type != expected_mime:
                logger.warning(f"MIME type mismatch: expected {expected_mime}, got {mime_type}")
                return False
            
            return True
            
        except Exception as e:
            logger.warning(f"MIME validation failed: {str(e)}")
            return True  # Continue processing even if validation fails
    
    async def _parse_file(self, file_path: str, source_url: str) -> Dict[str, Any]:
        """
        Parse a file asynchronously.
        
        Args:
            file_path: Path to the file to parse.
            source_url: Original URL of the file.
            
        Returns:
            Dictionary with parsed text and metadata.
        """
        # Run sync parser in thread pool to avoid blocking
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(
            None, self._parse_file_sync, Path(file_path), source_url
        )
    
    def _parse_file_sync(self, file_path: Path, source_url: str) -> Dict[str, Any]:
        """
        Parse a file synchronously.
        
        Args:
            file_path: Path to the file to parse.
            source_url: Original URL of the file.
            
        Returns:
            Dictionary with parsed text and metadata.
        """
        try:
            # Get file extension
            extension = file_path.suffix.lower()
            
            if extension not in self.supported_extensions:
                raise ParserError(f"Unsupported file type: {extension}")
            
            # Validate MIME type
            if not self._validate_mime_type(file_path, extension):
                logger.warning(f"MIME type validation failed for {file_path}")
            
            # Parse based on file type
            if extension == '.pdf':
                text, metadata = self._parse_pdf(file_path)
            elif extension == '.docx':
                text, metadata = self._parse_docx(file_path)
            elif extension == '.txt':
                text, metadata = self._parse_txt(file_path)
            elif extension == '.eml':
                text, metadata = self._parse_eml(file_path)
            else:
                raise ParserError(f"Unsupported file type: {extension}")
            
            # Clean and validate text
            text = self._clean_text(text)
            if not text.strip():
                raise ParserError("No text content extracted from file")
            
            # Create result structure
            result = {
                "filetype": extension[1:],  # Remove leading dot
                "url": source_url,
                "full_text": text,
                "length": len(text),
                "metadata": metadata or {}
            }
            
            return result
            
        except Exception as e:
            logger.error(f"Failed to parse file {file_path}: {str(e)}")
            raise ParserError(f"Parsing failed: {str(e)}")
    
    def _parse_pdf(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """
        Parse PDF file using PyMuPDF with OCR fallback.
        
        Args:
            file_path: Path to the PDF file.
            
        Returns:
            Tuple of (extracted text, metadata).
        """
        try:
            doc = fitz.open(str(file_path))
            text_parts = []
            metadata = {}
            
            # Extract PDF metadata
            if doc.metadata:
                metadata = {
                    'title': doc.metadata.get('title', ''),
                    'author': doc.metadata.get('author', ''),
                    'subject': doc.metadata.get('subject', ''),
                    'creator': doc.metadata.get('creator', ''),
                    'producer': doc.metadata.get('producer', ''),
                    'creationDate': doc.metadata.get('creationDate', ''),
                    'modDate': doc.metadata.get('modDate', ''),
                    'page_count': len(doc)
                }
            
            # Extract text from each page
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text = page.get_text()
                if text.strip():
                    text_parts.append(text)
            
            doc.close()
            
            # If no text extracted and OCR is enabled, try OCR
            if not text_parts and self.enable_ocr:
                logger.info(f"Attempting OCR for scanned PDF: {file_path}")
                text_parts = self._ocr_pdf_pages(file_path)
                if text_parts:
                    metadata['ocr_used'] = True
            
            if not text_parts:
                raise ParserError("PDF appears to be scanned - OCR required but not enabled")
            
            return "\n".join(text_parts), metadata
            
        except Exception as e:
            raise ParserError(f"PDF parsing failed: {str(e)}")
    
    def _ocr_pdf_pages(self, file_path: Path) -> List[str]:
        """
        Perform OCR on PDF pages.
        
        Args:
            file_path: Path to the PDF file.
            
        Returns:
            List of extracted text from each page.
        """
        if not OCR_AVAILABLE:
            raise ParserError("OCR not available - pytesseract and Pillow required")
        
        try:
            doc = fitz.open(str(file_path))
            text_parts = []
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                
                # Convert page to image
                mat = fitz.Matrix(2, 2)  # 2x zoom for better OCR
                pix = page.get_pixmap(matrix=mat)
                img_data = pix.tobytes("png")
                
                # Convert to PIL Image
                import io
                img = Image.open(io.BytesIO(img_data))
                
                # Perform OCR
                text = pytesseract.image_to_string(img)
                if text.strip():
                    text_parts.append(text)
            
            doc.close()
            return text_parts
            
        except Exception as e:
            logger.error(f"OCR failed: {str(e)}")
            return []
    
    def _parse_docx(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """
        Parse DOCX file using python-docx with table preservation.
        
        Args:
            file_path: Path to the DOCX file.
            
        Returns:
            Tuple of (extracted text, metadata).
        """
        try:
            doc = DocxDocument(str(file_path))
            text_parts = []
            tables_data = []
            metadata = {
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables),
                'section_count': len(doc.sections)
            }
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables with structure preservation
            for table_idx, table in enumerate(doc.tables):
                table_data = {
                    'table_index': table_idx,
                    'rows': [],
                    'headers': []
                }
                
                table_text_parts = []
                
                for row_idx, row in enumerate(table.rows):
                    row_data = []
                    row_text_parts = []
                    
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_data.append(cell_text)
                            row_text_parts.append(cell_text)
                    
                    if row_data:
                        table_data['rows'].append(row_data)
                        table_text_parts.append(' | '.join(row_text_parts))
                        
                        # First row is typically headers
                        if row_idx == 0:
                            table_data['headers'] = row_data
                
                if table_text_parts:
                    text_parts.extend(table_text_parts)
                    tables_data.append(table_data)
            
            metadata['tables'] = tables_data
            
            return "\n".join(text_parts), metadata
            
        except Exception as e:
            raise ParserError(f"DOCX parsing failed: {str(e)}")
    
    def _parse_txt(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """
        Parse TXT file using standard file reading.
        
        Args:
            file_path: Path to the TXT file.
            
        Returns:
            Tuple of (extracted text, metadata).
        """
        try:
            # Try UTF-8 first
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
                    encoding = 'utf-8'
            except UnicodeDecodeError:
                # Fallback to other encodings
                encodings = ['latin-1', 'cp1252', 'iso-8859-1']
                for encoding in encodings:
                    try:
                        with open(file_path, 'r', encoding=encoding) as f:
                            text = f.read()
                            break
                    except UnicodeDecodeError:
                        continue
                else:
                    raise ParserError("Unable to decode text file with any supported encoding")
            
            metadata = {
                'encoding': encoding,
                'line_count': len(text.splitlines()),
                'word_count': len(text.split())
            }
            
            return text, metadata
                
        except Exception as e:
            raise ParserError(f"TXT parsing failed: {str(e)}")
    
    def _parse_eml(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """
        Parse EML file using email parser with HTML support.
        
        Args:
            file_path: Path to the EML file.
            
        Returns:
            Tuple of (extracted text, metadata).
        """
        try:
            with open(file_path, 'rb') as f:
                msg = BytesParser(policy=policy.default).parse(f)
            
            text_parts = []
            metadata = {
                'subject': msg.get('subject', ''),
                'from': msg.get('from', ''),
                'to': msg.get('to', ''),
                'date': msg.get('date', ''),
                'content_type': msg.get_content_type(),
                'has_attachments': False
            }
            
            # Extract subject
            if metadata['subject']:
                text_parts.append(f"Subject: {metadata['subject']}")
            
            # Extract body content
            if msg.is_multipart():
                # Handle multipart messages
                for part in msg.walk():
                    content_type = part.get_content_type()
                    
                    if content_type == "text/plain":
                        content = part.get_content()
                        if content:
                            text_parts.append(content)
                    elif content_type == "text/html" and HTML2TEXT_AVAILABLE:
                        # Convert HTML to text
                        html_content = part.get_content()
                        if html_content:
                            h = html2text.HTML2Text()
                            h.ignore_links = True
                            h.ignore_images = True
                            text_content = h.handle(html_content)
                            if text_content.strip():
                                text_parts.append(text_content)
                    elif part.get_filename():
                        metadata['has_attachments'] = True
            
            else:
                # Handle single part messages
                content_type = msg.get_content_type()
                if content_type == "text/plain":
                    content = msg.get_content()
                    if content:
                        text_parts.append(content)
                elif content_type == "text/html" and HTML2TEXT_AVAILABLE:
                    # Convert HTML to text
                    html_content = msg.get_content()
                    if html_content:
                        h = html2text.HTML2Text()
                        h.ignore_links = True
                        h.ignore_images = True
                        text_content = h.handle(html_content)
                        if text_content.strip():
                            text_parts.append(text_content)
            
            if not text_parts:
                return "No text content found", metadata
            
            return "\n".join(text_parts), metadata
            
        except Exception as e:
            raise ParserError(f"EML parsing failed: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """
        Clean and normalize extracted text.
        
        Args:
            text: Raw extracted text.
            
        Returns:
            Cleaned text.
        """
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = ' '.join(text.split())
        
        # Remove common headers/footers (basic implementation)
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            # Skip common header/footer patterns
            if (line and 
                not line.isdigit() and  # Skip page numbers
                len(line) > 3 and  # Skip very short lines
                not line.startswith('Page ') and
                not line.startswith('Confidential') and
                not line.startswith('Draft')):
                cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines)
    
    def _cleanup_temp_file(self, temp_path: str) -> None:
        """
        Clean up temporary file.
        
        Args:
            temp_path: Path to the temporary file to delete.
        """
        try:
            if os.path.exists(temp_path):
                os.unlink(temp_path)
                logger.debug(f"Cleaned up temporary file: {temp_path}")
        except Exception as e:
            logger.warning(f"Failed to cleanup temp file {temp_path}: {str(e)}")


# Convenience functions for easy usage
async def parse_url(url: str, enable_ocr: bool = False) -> Dict[str, Any]:
    """
    Parse a file from URL.
    
    Args:
        url: The URL pointing to the file to parse.
        enable_ocr: Whether to enable OCR fallback for scanned PDFs.
        
    Returns:
        Dictionary containing parsed text and metadata.
    """
    parser = FileParser(enable_ocr=enable_ocr)
    return await parser.parse_from_url(url)


def parse_file(file_path: str, enable_ocr: bool = False) -> Dict[str, Any]:
    """
    Parse a file from local path.
    
    Args:
        file_path: Path to the local file to parse.
        enable_ocr: Whether to enable OCR fallback for scanned PDFs.
        
    Returns:
        Dictionary containing parsed text and metadata.
    """
    parser = FileParser(enable_ocr=enable_ocr)
    return parser.parse_from_file(file_path)


if __name__ == "__main__":
    # Example usage
    import asyncio
    
    async def test_parser():
        # Test with a sample URL (replace with actual URL)
        try:
            result = await parse_url("https://example.com/sample.pdf", enable_ocr=True)
            print("Parsed result:", result)
        except ParserError as e:
            print(f"Parser error: {e}")
    
    # Run test
    asyncio.run(test_parser())
